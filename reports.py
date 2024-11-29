import datetime
import sqlite3
import csv
from datetime import datetime
from docx import Document  # Импортируем библиотеку для работы с документами Word
from func_with_time import calculate_time_difference, time_now  # Импортируем функции для работы с временем


def create_report_of_day():
    # Функция для создания отчета за день
    document = Document()  # Создание нового документа Word

    # Инициализация счетчиков для различных метрик
    count_of_applications = 0  # Общее количество закрытых заявок
    equipment_under_repair = 0  # Количество оборудования в ремонте
    unused_equipment = 0  # Общее время простаивания оборудования
    count_applications = 0  # Общее количество заявок

    # Подключение к базе данных
    con = sqlite3.connect("db/tab.db")
    cur = con.cursor()

    # Запрос всех записей из таблицы repair_hardware
    req = cur.execute("""select * from repair_hardware""").fetchall()
    for i in req:
        start = i[2]  # Время начала ремонта
        end = i[3]  # Время окончания ремонта
        done = i[7]  # Статус завершения ремонта
        count_of_applications += done  # Увеличиваем счетчик закрытых заявок
        if done == 0:
            equipment_under_repair += 1  # Увеличиваем количество оборудования, находящегося в ремонте

        print(end)  # Выводим дату окончания ремонта
        if end is None:
            # Если время окончания отсутствует, вычисляем время простаивания оборудования
            unused_equipment += calculate_time_difference(start, time_now())

        count_applications += 1  # Увеличиваем общее количество заявок

    # Добавление заголовков и информации в документ
    document.add_heading('Отчет Бригады', 0)
    p = document.add_paragraph(f'ФИО работников: ')  # Добавляем информацию о работниках
    document.add_paragraph(f'Общее количество заявок: {count_applications}')  # Общее количество заявок
    document.add_paragraph(f'Количество закрытых заявок: {count_of_applications}')  # Количество закрытых заявок
    document.add_paragraph(f'Количество оборудования в ремонте: {equipment_under_repair}')  # Оборудование в ремонте
    document.add_paragraph(
        f"Оборудование простаивало(в часах): {unused_equipment}")  # Общее время простаивания оборудования
    document.save("reports/test.doc")  # Сохранение документа с отчетом в указанной директории


def benefits_calculating():
    # Функция для расчета выгод от ремонта оборудования
    result = 0  # Инициализация переменной для хранения результата
    con = sqlite3.connect("db/tab.db")  # Подключение к базе данных
    cur = con.cursor()

    # Запрос деталей оборудования и времени начала и окончания ремонта
    req = cur.execute("""select hardware.details, repair_hardware.start, repair_hardware.end from repair_hardware
    inner join hardware on repair_hardware.id_hardware = hardware.id""").fetchall()

    # Проход по каждой записи и вычисление результата
    for i in req:
        # Расчет выгод и накопление результата
        result += (i[0] / 24) * calculate_time_difference(i[1], i[2])  # Объем оборудования * время ремонта в днях

    return round(result)  # Возвращение результата, округленного до целого


# Печатаем результат расчета выгод
print(benefits_calculating())
create_report_of_day()  # Вызов функции для создания отчета за день